from docx import Document
from fpdf import FPDF
import pdfplumber
import os

class FileProcessor:
    def extract_text_from_pdf(self, file_path: str) -> str:
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return '\n'.join(text)

    def extract_text_from_word(self, file_path: str) -> str:
        doc = Document(file_path)
        paras = [p.text for p in doc.paragraphs]
        return '\n'.join(paras)

    def extract_text_from_markdown(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def save_text_to_word(self, text: str, out_path: str):
        doc = Document()
        for para in text.split('\n'):
            doc.add_paragraph(para)
        doc.save(out_path)

    def save_text_to_pdf(self, text: str, out_path: str):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        font_dir = os.path.join(os.path.dirname(__file__), '../static/fonts')
        font_path = os.path.join(font_dir, 'msyh.ttc')
        if not os.path.exists(font_path):
            raise RuntimeError('缺少字体文件 msyh.ttc，请将微软雅黑字体(msyh.ttc)放到 static/fonts/ 目录下')
        pdf.add_font('MSYH', '', font_path, uni=True)
        pdf.set_font('MSYH', size=12)
        for line in text.split('\n'):
            pdf.multi_cell(0, 10, line)
        pdf.output(out_path)

    def save_text_to_markdown(self, text: str, out_path: str):
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(text)
